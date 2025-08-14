#!/usr/bin/env python3
"""Demo script to showcase multi-format document analysis."""

import sys
import os
import tempfile
from pathlib import Path

# Add the app directory to Python path
app_dir = Path(__file__).parent / "app"
sys.path.insert(0, str(app_dir))

def create_comprehensive_docx():
    """Create a comprehensive DOCX document for testing."""
    try:
        from docx import Document
        
        doc = Document()
        doc.add_heading('EMPLOYMENT AGREEMENT', 0)
        
        doc.add_paragraph('This Employment Agreement ("Agreement") is entered into on January 15, 2024, between TechCorp Inc., a Delaware corporation ("Company"), and Sarah Johnson ("Employee").')
        
        doc.add_heading('1. POSITION AND DUTIES', level=1)
        doc.add_paragraph('Employee shall serve as Senior Software Engineer and shall perform such duties as are customarily associated with such position.')
        
        doc.add_heading('2. COMPENSATION', level=1)
        doc.add_paragraph('Company shall pay Employee an annual salary of $95,000.00, payable in accordance with Company\'s standard payroll practices.')
        
        doc.add_heading('3. BENEFITS', level=1)
        doc.add_paragraph('Employee shall be entitled to participate in all employee benefit plans maintained by Company, including health insurance, dental coverage, and 401(k) retirement plan with 4% company matching.')
        
        doc.add_heading('4. CONFIDENTIALITY', level=1)
        doc.add_paragraph('Employee agrees to maintain the confidentiality of all proprietary information and trade secrets of the Company during and after employment.')
        
        doc.add_heading('5. TERMINATION', level=1)
        doc.add_paragraph('Either party may terminate this Agreement at any time with or without cause by providing thirty (30) days written notice to the other party.')
        
        doc.add_paragraph('IN WITNESS WHEREOF, the parties have executed this Agreement as of the date first written above.')
        
        # Save to temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        return temp_file.name
        
    except Exception as e:
        print(f"❌ Failed to create comprehensive DOCX: {e}")
        return None

def create_nda_txt():
    """Create a comprehensive NDA text file."""
    content = """NON-DISCLOSURE AGREEMENT

This Non-Disclosure Agreement ("Agreement") is entered into on March 10, 2024, between:

DISCLOSING PARTY: InnovateTech Solutions LLC
Address: 123 Innovation Drive, Tech City, CA 90210
Email: legal@innovatetech.com

RECEIVING PARTY: Strategic Partners Inc.
Address: 456 Business Blvd, Commerce City, NY 10001
Email: contracts@strategicpartners.com

RECITALS

WHEREAS, Disclosing Party possesses certain confidential and proprietary information relating to advanced AI algorithms and machine learning models;

WHEREAS, Receiving Party desires to evaluate potential business opportunities that may require access to such confidential information;

NOW, THEREFORE, in consideration of the mutual covenants contained herein, the parties agree as follows:

1. DEFINITION OF CONFIDENTIAL INFORMATION

"Confidential Information" shall include all technical data, trade secrets, know-how, research, product plans, products, services, customers, customer lists, markets, software, developments, inventions, processes, formulas, technology, designs, drawings, engineering, hardware configuration information, marketing, finances, or other business information disclosed by Disclosing Party.

2. OBLIGATIONS OF RECEIVING PARTY

Receiving Party agrees to:
a) Hold and maintain all Confidential Information in strict confidence
b) Not disclose any Confidential Information to third parties without prior written consent
c) Use Confidential Information solely for evaluation purposes
d) Return or destroy all Confidential Information upon request

3. TERM

This Agreement shall remain in effect for a period of five (5) years from the date of execution, unless terminated earlier by mutual written consent.

4. REMEDIES

Receiving Party acknowledges that any breach of this Agreement may cause irreparable harm to Disclosing Party, and that monetary damages may be inadequate. Therefore, Disclosing Party shall be entitled to seek injunctive relief and other equitable remedies.

5. GOVERNING LAW

This Agreement shall be governed by and construed in accordance with the laws of the State of California.

IN WITNESS WHEREOF, the parties have executed this Agreement as of the date first written above.

DISCLOSING PARTY:                    RECEIVING PARTY:
InnovateTech Solutions LLC           Strategic Partners Inc.

By: _________________________       By: _________________________
Name: Michael Chen                   Name: Jennifer Rodriguez
Title: CEO                          Title: VP Business Development
Date: March 10, 2024                Date: March 10, 2024
"""
    
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.txt', mode='w')
    temp_file.write(content)
    temp_file.close()
    return temp_file.name

def analyze_document_comprehensive(file_path, file_type):
    """Perform comprehensive analysis and display results."""
    print(f"\n📄 Analyzing {file_type.upper()} Document: {os.path.basename(file_path)}")
    print("=" * 60)
    
    try:
        from streamlit_app import StandaloneLegalAnalyzer
        analyzer = StandaloneLegalAnalyzer()
        
        # Perform analysis
        result = analyzer.analyze_document(file_path, max_clauses=10)
        
        # Display results
        print(f"✅ Analysis completed successfully!")
        print(f"\n📋 DOCUMENT OVERVIEW:")
        print(f"   • File Type: {result['document_info']['file_type'].upper()}")
        print(f"   • Document Type: {result['classification']['predicted_type'].replace('_', ' ').title()}")
        print(f"   • Confidence: {result['classification']['confidence']:.1%}")
        print(f"   • Word Count: {result['text_statistics']['word_count']:,}")
        print(f"   • Character Count: {result['text_statistics']['character_count']:,}")
        print(f"   • Sentence Count: {result['text_statistics']['sentence_count']}")
        print(f"   • Clauses Found: {result['clauses']['total_count']}")
        
        # Display entities
        print(f"\n🏷️ EXTRACTED ENTITIES:")
        entities = result['entities']
        
        if entities.get('key_parties'):
            print(f"   • Key Parties: {', '.join([p['name'] for p in entities['key_parties'][:3]])}")
        
        for entity_type, entity_list in entities['entities'].items():
            if entity_list:
                values = [e['text'] for e in entity_list[:3]]
                print(f"   • {entity_type.replace('_', ' ').title()}: {', '.join(values)}")
        
        # Display summary
        print(f"\n📝 SUMMARY:")
        summary = result.get('summary', {})
        if summary.get('key_findings'):
            for finding in summary['key_findings']:
                print(f"   • {finding}")
        
        # Display recommendations
        print(f"\n💡 RECOMMENDATIONS:")
        for rec in result.get('recommendations', []):
            print(f"   • {rec}")
        
        # Display simplified clauses
        if result['clauses']['simplified_clauses']:
            print(f"\n📋 SIMPLIFIED CLAUSES (showing first 3):")
            for i, clause_data in enumerate(result['clauses']['simplified_clauses'][:3]):
                print(f"\n   Clause {i+1}:")
                print(f"   📜 Original: {clause_data['original_clause'][:100]}...")
                print(f"   ✨ Simplified: {clause_data['simplified_clause'][:100]}...")
                print(f"   🎯 Summary: {clause_data['plain_english_summary']}")
        
        return True
        
    except Exception as e:
        print(f"❌ Analysis failed: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    """Run comprehensive multi-format demo."""
    print("🚀 LEGAL DOCUMENT ANALYZER - MULTI-FORMAT DEMO")
    print("=" * 60)
    
    # Test 1: Analyze existing TXT sample
    print("\n🔍 TEST 1: Analyzing existing NDA sample (TXT)")
    txt_sample = "sample_documents/sample_nda.txt"
    if os.path.exists(txt_sample):
        analyze_document_comprehensive(txt_sample, "txt")
    else:
        print("⚠️ Sample NDA file not found")
    
    # Test 2: Create and analyze comprehensive DOCX
    print("\n🔍 TEST 2: Creating and analyzing Employment Agreement (DOCX)")
    docx_file = create_comprehensive_docx()
    if docx_file:
        success = analyze_document_comprehensive(docx_file, "docx")
        os.unlink(docx_file)  # Clean up
    
    # Test 3: Create and analyze comprehensive NDA TXT
    print("\n🔍 TEST 3: Creating and analyzing comprehensive NDA (TXT)")
    nda_file = create_nda_txt()
    if nda_file:
        success = analyze_document_comprehensive(nda_file, "txt")
        os.unlink(nda_file)  # Clean up
    
    # Test 4: Show format support
    print("\n🔧 SUPPORTED FORMATS:")
    print("=" * 30)
    from streamlit_app import StandaloneLegalAnalyzer
    analyzer = StandaloneLegalAnalyzer()
    
    for fmt, available in analyzer.supported_formats.items():
        status = "✅ Available" if available else "❌ Not Available"
        print(f"   {fmt.upper()}: {status}")
    
    print("\n🎉 DEMO COMPLETED!")
    print("=" * 60)
    print("The Legal Document Analyzer supports multiple formats and provides:")
    print("• Document classification with confidence scores")
    print("• Comprehensive entity extraction (parties, dates, money, etc.)")
    print("• Intelligent clause extraction and simplification")
    print("• Plain English summaries")
    print("• Actionable recommendations")
    print("• Multi-format support (TXT, DOCX, PDF)")

if __name__ == "__main__":
    main()
