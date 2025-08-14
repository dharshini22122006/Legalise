#!/usr/bin/env python3
"""Test multi-format document support."""

import sys
import os
import tempfile
from pathlib import Path

# Add the app directory to Python path
app_dir = Path(__file__).parent / "app"
sys.path.insert(0, str(app_dir))

def create_sample_docx():
    """Create a sample DOCX file for testing."""
    try:
        from docx import Document
        
        doc = Document()
        doc.add_heading('Service Agreement', 0)
        doc.add_paragraph('This Service Agreement is entered into between Client Corp and Service Provider LLC.')
        doc.add_paragraph('The service provider agrees to deliver consulting services for a fee of $5,000.')
        doc.add_paragraph('The agreement shall remain in effect for 12 months from the effective date.')
        doc.add_paragraph('Either party may terminate this agreement with 30 days written notice.')
        
        # Save to temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        return temp_file.name
        
    except Exception as e:
        print(f"❌ Failed to create DOCX file: {e}")
        return None

def create_sample_pdf():
    """Create a sample PDF file for testing."""
    try:
        # For this test, we'll create a simple text file and rename it
        # In a real scenario, you'd use a PDF library like reportlab
        content = """Employment Contract

This Employment Contract is between ABC Company and John Doe.

Position: Software Developer
Salary: $75,000 per year
Start Date: January 1, 2024

Benefits include health insurance and 401k matching.

The employee agrees to maintain confidentiality of company information.

This contract may be terminated by either party with two weeks notice.
"""
        
        # For testing purposes, we'll just create a text file
        # The app should handle this gracefully
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.txt', mode='w')
        temp_file.write(content)
        temp_file.close()
        
        # Rename to .pdf for testing (the app will try to parse it as PDF and may fail gracefully)
        pdf_path = temp_file.name.replace('.txt', '.pdf')
        os.rename(temp_file.name, pdf_path)
        return pdf_path
        
    except Exception as e:
        print(f"❌ Failed to create PDF file: {e}")
        return None

def test_format_support():
    """Test multi-format document support."""
    print("🔍 Testing Multi-Format Document Support")
    print("=" * 50)
    
    from streamlit_app import StandaloneLegalAnalyzer
    analyzer = StandaloneLegalAnalyzer()
    
    # Test TXT format
    print("\n📄 Testing TXT format...")
    try:
        txt_file = "sample_documents/sample_nda.txt"
        if os.path.exists(txt_file):
            result = analyzer.analyze_document(txt_file, max_clauses=3)
            print(f"✅ TXT analysis successful")
            print(f"   - Type: {result['classification']['predicted_type']}")
            print(f"   - Words: {result['text_statistics']['word_count']}")
        else:
            print(f"⚠️ TXT sample file not found")
    except Exception as e:
        print(f"❌ TXT analysis failed: {e}")
    
    # Test DOCX format
    print("\n📄 Testing DOCX format...")
    docx_file = create_sample_docx()
    if docx_file:
        try:
            result = analyzer.analyze_document(docx_file, max_clauses=3)
            print(f"✅ DOCX analysis successful")
            print(f"   - Type: {result['classification']['predicted_type']}")
            print(f"   - Words: {result['text_statistics']['word_count']}")
            os.unlink(docx_file)  # Clean up
        except Exception as e:
            print(f"❌ DOCX analysis failed: {e}")
            if os.path.exists(docx_file):
                os.unlink(docx_file)
    
    # Test PDF format (this might fail since we're creating a fake PDF)
    print("\n📄 Testing PDF format...")
    pdf_file = create_sample_pdf()
    if pdf_file:
        try:
            result = analyzer.analyze_document(pdf_file, max_clauses=3)
            print(f"✅ PDF analysis successful")
            print(f"   - Type: {result['classification']['predicted_type']}")
            print(f"   - Words: {result['text_statistics']['word_count']}")
        except Exception as e:
            print(f"⚠️ PDF analysis failed (expected for fake PDF): {e}")
        finally:
            if os.path.exists(pdf_file):
                os.unlink(pdf_file)
    
    # Test supported formats detection
    print("\n🔧 Testing format support detection...")
    supported = analyzer.supported_formats
    print("Supported formats:")
    for fmt, available in supported.items():
        status = "✅" if available else "❌"
        print(f"   {status} {fmt.upper()}")
    
    print("\n" + "=" * 50)
    print("🎉 Multi-format testing completed!")

def test_entity_extraction():
    """Test entity extraction capabilities."""
    print("\n🏷️ Testing Entity Extraction")
    print("=" * 30)
    
    from streamlit_app import StandaloneLegalAnalyzer
    analyzer = StandaloneLegalAnalyzer()
    
    test_text = """
    This Service Agreement is between Acme Corp and John Smith.
    The total contract value is $50,000.00 and payment is due by December 31, 2024.
    Contact information: john.smith@email.com or call 555-123-4567.
    The service fee is 15% of the total project cost.
    """
    
    entities = analyzer._extract_entities(test_text)
    
    print("Extracted entities:")
    for entity_type, entity_list in entities['entities'].items():
        if entity_list:
            print(f"  {entity_type}: {[e['text'] for e in entity_list]}")
    
    if entities['key_parties']:
        print(f"  Key parties: {[p['name'] for p in entities['key_parties']]}")

if __name__ == "__main__":
    test_format_support()
    test_entity_extraction()
